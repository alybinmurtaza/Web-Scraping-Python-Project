# export_task1_reports.py
import os
from datetime import datetime
from urllib.parse import urlparse

import pandas as pd



# - Helpers -
def now_iso():
    # Local time with timezone offset, second precision
    return datetime.now().astimezone().isoformat(timespec="seconds")

def safe_filename(stem: str, suffix: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{stem}_{ts}.{suffix.lstrip('.')}"

# - Load Data -
CSV_PATH = "dawn_top30_headlines.csv"
if os.path.exists(CSV_PATH):
    df = pd.read_csv(CSV_PATH)
else:
    # Fallback: call your function if CSV missing
    # df = collect_dawn_headlines(30)
    raise FileNotFoundError(
        "dawn_top30_headlines.csv not found."
    )

# Here I am ensuring expected columns exist for transparency
for col in ["headline", "url", "scraped_at", "source"]:
    if col not in df.columns:

        if col == "scraped_at":
            df[col] = now_iso()
        elif col == "source":
            df[col] = "dawn.com"
        else:
            df[col] = ""

# - Build .DOCX -
from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):

    #  Here I am handling "No URL available"
    if not url or not url.startswith(("http://", "https://")):
        paragraph.add_run(text or "No URL available")
        return


    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)


    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)


    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    rPr.append(u); rPr.append(color)
    new_run.append(rPr)


    t = OxmlElement("w:t"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def export_docx(df: pd.DataFrame, path: str):
    doc = Document()
    doc.core_properties.title = "CS4048 – Task 1: Dawn News Headlines"
    doc.core_properties.subject = "Data scraping transparency report"
    doc.core_properties.author = "Student – CS4048"
    doc.core_properties.created = datetime.now()

    # Title + metadata
    doc.add_heading("CS4048 – Task 1: Dawn News Headlines", level=0)
    doc.add_paragraph(f"Generated: {now_iso()}")
    doc.add_paragraph(f"Rows: {len(df)} | Source: dawn.com")

    doc.add_paragraph(
        "This table includes the headline text, original article URL, and the scraper timestamp for reproducibility."
    )

    # Table
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Headline"
    hdr[2].text = "URL"
    hdr[3].text = "Scraped At"
    hdr[4].text = "Source"

    for i, row in df.reset_index(drop=True).iterrows():
        r = table.add_row().cells
        r[0].text = str(i + 1)
        r[1].text = str(row.get("headline", ""))


        p = r[2].paragraphs[0]
        url_text = urlparse(str(row.get("url", ""))).netloc or str(row.get("url", ""))
        add_hyperlink(p, url_text or "No URL available", str(row.get("url", "")))

        r[3].text = str(row.get("scraped_at", ""))
        r[4].text = str(row.get("source", ""))


    try:
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(4.0)
        table.columns[2].width = Inches(2.0)
        table.columns[3].width = Inches(1.6)
        table.columns[4].width = Inches(1.2)
    except Exception:
        pass

    doc.save(path)
    return path

# - Build .PDF -
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet

def export_pdf(df: pd.DataFrame, path: str):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=36, rightMargin=36, topMargin=48, bottomMargin=48
    )
    story = []
    story.append(Paragraph("CS4048 – Task 1: Dawn News Headlines", styles["Title"]))
    story.append(Paragraph(f"Generated: {now_iso()}", styles["Normal"]))
    story.append(Paragraph(f"Rows: {len(df)} | Source: dawn.com", styles["Normal"]))
    story.append(Spacer(1, 12))


    data = [["#", "Headline", "URL"]]
    for i, row in df.reset_index(drop=True).iterrows():
        head = Paragraph(str(row.get("headline", "")), styles["BodyText"])
        url_val = str(row.get("url", "No URL available"))
        if url_val.startswith(("http://", "https://")):
            url_par = Paragraph(f'<link href="{url_val}">{url_val}</link>', styles["BodyText"])
        else:
            url_par = Paragraph(url_val, styles["BodyText"])
        data.append([str(i + 1), head, url_par])

    table = Table(data, colWidths=[25, 270, 156])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)

    doc.build(story)
    return path


docx_out = safe_filename("Task1_Dawn_Report", "docx")
pdf_out = safe_filename("Task1_Dawn_Report", "pdf")

export_docx(df, docx_out)
export_pdf(df, pdf_out)

print(f"Saved:\n- {docx_out}\n- {pdf_out}")
