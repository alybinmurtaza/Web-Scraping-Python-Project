

import os
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors


# -------------------- Small utils --------------------

def now_iso() -> str:
    """Local timezone ISO string with seconds."""
    return datetime.now().astimezone().isoformat(timespec="seconds")


def safe_filename(stem: str, suffix: str) -> str:
    """Make a timestamped file name like 'stem_20250906_181530.pdf'."""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{stem}_{ts}.{suffix.lstrip('.')}"


def ensure_columns(df: pd.DataFrame, columns):
    """Ensure all columns exist (adds missing as blank) and return reordered copy."""
    out = df.copy()
    for c in columns:
        if c not in out.columns:
            out[c] = ""
    return out[columns]


# -------------------- DOCX export (generic table) --------------------

def export_docx_table(df: pd.DataFrame, path: str, title: str, subtitle: str = "", columns=None):
    """
    Create a .docx with a table from an arbitrary DataFrame.
    - df: pandas DataFrame
    - path: output filename like "Task2_PSX_Indices_Report.docx"
    - title/subtitle: free text (subtitle can include “As of …” or notes)
    - columns: optional list to reorder/limit columns
    """
    use_df = df.copy()
    if columns:
        use_df = ensure_columns(use_df, columns)
    else:
        columns = list(use_df.columns)

    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = "Student – CS4048"
    doc.core_properties.created = datetime.now()

    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {now_iso()}")
    if subtitle:
        doc.add_paragraph(subtitle)

    table = doc.add_table(rows=1, cols=len(columns))
    hdr = table.rows[0].cells
    for i, c in enumerate(columns):
        hdr[i].text = str(c)

    # rows
    for _, row in use_df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(columns):
            val = row[c]
            cells[i].text = "" if (pd.isna(val) or val is None) else str(val)

    # best-effort column sizing (won't always apply in Word UI, but helps)
    try:
        # If there are many columns, set a narrow default width
        for col in table.columns:
            col.width = Inches(1.3)
    except Exception:
        pass

    doc.save(path)
    return path


# -------------------- PDF export (generic table) --------------------

def export_pdf_table(df: pd.DataFrame, path: str, title: str, columns=None):
    """
    Create a .pdf with a table from an arbitrary DataFrame.
    - columns: optional list to reorder/limit columns
    """
    use_df = df.copy()
    if columns:
        use_df = ensure_columns(use_df, columns)
    else:
        columns = list(use_df.columns)

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=36, rightMargin=36, topMargin=48, bottomMargin=48
    )
    story = []
    story.append(Paragraph(title, styles["Title"]))
    story.append(Paragraph(f"Generated: {now_iso()}", styles["Normal"]))
    story.append(Spacer(1, 10))

    # Build table data
    data = [columns]
    for _, row in use_df.iterrows():
        data.append([("" if (pd.isna(row[c]) or row[c] is None) else str(row[c])) for c in columns])

    # Simple table (ReportLab handles page breaks automatically)
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)

    doc.build(story)
    return path


# -------------------- Main runner --------------------

def main():
    # ---------- Indices (required) ----------
    indices_csv = "psx_indices.csv"
    if not os.path.exists(indices_csv):
        raise FileNotFoundError(
            f"'{indices_csv}' not found. Run your Task-2 scraper first to create it."
        )
    dfi = pd.read_csv(indices_csv)

    # Preferred column order for indices (keep NA fields visible for transparency)
    cols_idx = [
        "Index Name", "High", "Low", "Current", "Change", "Change %",
        "LDCP", "Open", "Volume", "scraped_at", "source"
    ]
    subtitle_idx = "As provided by PSX Data Portal (Indices)."
    if "source_as_of" in dfi.columns:
        # Add PSX's own timestamp if the scraper captured it
        cols_idx.append("source_as_of")
        try:
            asof = dfi["source_as_of"].iloc[0]
            subtitle_idx = f"As provided by PSX (source_as_of: {asof})"
        except Exception:
            pass

    # Create files
    docx_idx = safe_filename("Task2_PSX_Indices_Report", "docx")
    pdf_idx  = safe_filename("Task2_PSX_Indices_Report", "pdf")

    export_docx_table(dfi, docx_idx, title="CS4048 – Task 2: PSX Indices",
                      subtitle=subtitle_idx, columns=cols_idx)
    export_pdf_table(dfi,  pdf_idx,  title="CS4048 – Task 2: PSX Indices",
                     columns=cols_idx)

    print("Saved:")
    print(" -", docx_idx)
    print(" -", pdf_idx)

    # ---------- Main Board (optional) ----------
    mainboard_csv = "psx_mainboard.csv"
    if os.path.exists(mainboard_csv):
        dfm = pd.read_csv(mainboard_csv)
    else:
        dfm = pd.DataFrame()

    if len(dfm):
        cols_mb = ["Index Name", "LDCP", "Open", "High", "Low", "Current", "Change", "Volume", "scraped_at", "source"]

        docx_mb = safe_filename("Task2_PSX_MainBoard_Report", "docx")
        pdf_mb  = safe_filename("Task2_PSX_MainBoard_Report", "pdf")

        export_docx_table(dfm, docx_mb, title="CS4048 – Task 2: PSX Main Board", columns=cols_mb)
        export_pdf_table(dfm,  pdf_mb,  title="CS4048 – Task 2: PSX Main Board", columns=cols_mb)

        print(" -", docx_mb)
        print(" -", pdf_mb)
    else:
        print("No Main Board rows found (page likely JS-rendered). Skipping Main Board report.")


if __name__ == "__main__":
    main()
